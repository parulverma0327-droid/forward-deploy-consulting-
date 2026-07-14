"""Generate UC1 Agent Specs.docx from agent definitions."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "UC1 Agent Specs.docx"

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
TEAL = RGBColor(0x1A, 0x6B, 0x5A)
GREY = RGBColor(0x55, 0x55, 0x55)

AGENTS = [
    {
        "name": "ForecastSignalBQAgent",
        "role": (
            "BigQuery-native agent on Gold hubs. Detects scopes where consumer or forecast signals "
            "changed since the last agent run, writes signal_delta, and answers on-demand planner "
            "questions (Inquiry path) via NL→SQL. Does not score unit lifts, route specialists, "
            "or publish recommendations."
        ),
        "status": "Deployed in BigQuery — Agent catalog + scheduled query UC1 – refresh signal_delta",
        "inputs": [
            (
                "Clickstream Agg",
                "product_cd, region, week, weighted_intent_score, search_count, pdp_view_count, "
                "cart_add_count, last_updated_dt",
            ),
            (
                "DemandForecast_AsIS",
                "product_cd, region, week, units_forecast, last_updated_dt",
            ),
            ("Member Hub", "region, ds_model_refresh_dt, membership_tier"),
            ("agent_run_history", "run_id, completed_at, scope_json"),
            ("Geo Region", "enterprise_region, region_canonical"),
        ],
        "computation": [
            "Last-run boundary = MAX(completed_at) from agent_run_history; first run uses 24h lookback",
            "Intent delta — compare Clickstream Agg vs prior snapshot; emit when weighted_intent change ≥ 15%",
            "Baseline delta — compare DemandForecast_AsIS.units_forecast vs prior; emit when change ≥ 10%",
            "delta_type = intent_spike | intent_drop | baseline_shift | member_model_refresh",
            "Write rows to signal_delta with detected_at = CURRENT_TIMESTAMP()",
            "Inquiry path (on demand) — same BQ agent; planner NL → SQL on Gold",
        ],
        "output_fields": [
            ("signal_delta_id", "Row PK (UUID)", "signal_delta"),
            ("product_cd", "Style + color where signal moved", "signal_delta"),
            ("region", "Enterprise region from source hub", "signal_delta"),
            ("region_canonical", "Planning region (e.g. US-PNW)", "signal_delta"),
            ("week", "Affected forecast week", "signal_delta"),
            ("delta_type", "intent_spike | intent_drop | baseline_shift | member_model_refresh", "signal_delta"),
            ("delta_magnitude", "% or absolute change in triggering metric", "signal_delta"),
            ("source_table", "Hub that triggered row — Clickstream Agg, DemandForecast_AsIS", "signal_delta"),
            ("detected_at", "UTC when BQ agent wrote this row", "signal_delta"),
            ("prior_value", "Metric before change (audit)", "signal_delta"),
            ("current_value", "Metric after change (audit)", "signal_delta"),
        ],
        "does_not": [
            "Run gates or issue run_id (OrchestratorAgent)",
            "Compute intent_units_lift or replacement_units_lift",
            "Publish recommendations (PublisherAgent)",
        ],
    },
    {
        "name": "OrchestratorAgent",
        "role": (
            "Hourly traffic controller. Reads signal_delta from ForecastSignalBQAgent, runs freshness "
            "gates, builds scope_json, issues run_id, routes specialist agents "
            "(RC ∥ DS → IP → Traceability → Publisher). Does not detect changes or score forecasts."
        ),
        "status": "Agent — Cloud Run (agents/orchestrator/, POST /run)",
        "inputs": [
            (
                "signal_delta",
                "product_cd, region, region_canonical, week, delta_type, detected_at, source_table",
            ),
            ("Clickstream Agg", "last_updated_dt"),
            ("DemandForecast_AsIS", "last_updated_dt"),
            ("Member Hub", "ds_model_refresh_dt, region"),
        ],
        "computation": [
            "Bucket 1 gates — freshness < 2h on Clickstream Agg + DemandForecast_AsIS; member model current",
            "Bucket 2 — read signal_delta since last completed run; empty → skip",
            "Build scope_json from delta rows + Product sku lookup",
            "Issue run_id; write agent_run_history with run_status, scope_json, timestamps",
        ],
        "output_fields": [
            ("run_id", "Run PK — FK on all downstream agent tables", "agent_run_history"),
            ("run_status", "completed | skipped | failed", "agent_run_history"),
            ("gates_passed", "TRUE if Bucket 1 passed", "agent_run_history"),
            ("scope_json", "Routed scopes — product_cd, sku_id, region_canonical, week", "agent_run_history"),
            ("agents_executed", 'JSON list e.g. ["BQ","ORCH","RC","DS","IP","TR","PUB"]', "agent_run_history"),
            ("started_at", "Run start UTC", "agent_run_history"),
            ("completed_at", "Run end UTC", "agent_run_history"),
        ],
        "does_not": [
            "Detect signal changes on Gold (ForecastSignalBQAgent)",
            "Score intent, repurchase, or safety stock",
            "Write demand_forecast_recommendation or publish recommendations",
        ],
    },
    {
        "name": "DemandSensingAgent",
        "role": (
            "Reads aggregated clickstream intent and baseline forecast. Converts intent into unit "
            "lifts and proposes an enriched forecast for Run 2. Does not build baseline or auto-apply."
        ),
        "status": "Built — agents/demand_sensing/",
        "inputs": [
            (
                "demand_forecast_base",
                "product_cd, sku_id, region, week, week_start_dt, season_code, category, "
                "units_forecast, ly_same_week_sales, promo_factor",
            ),
            (
                "clickstream_agg",
                "product_cd, region, week, week_start_dt, search_count, pdp_view_count, "
                "wishlist_add_count, cart_add_count, weighted_intent_score, member_intent_count, "
                "guest_intent_count, top_search_query, browse_score",
            ),
            (
                "member_hub",
                "member_id, region, membership_tier, clv_tier, visit_count_12m, order_count_12m",
            ),
            ("geo_region", "enterprise_region, region_canonical, source_system"),
            ("product_global_line_plan", "product_code, style_name"),
        ],
        "input_note": (
            "Upstream: clickstream_base feeds clickstream_agg. Fan-in: reads replacement_score "
            "after ReplacementCycleAgent completes."
        ),
        "computation": [
            "weighted_intent_score = Σ(event_count × intent_weight) — search=0.2, pdp_view=0.4, wishlist=0.8, cart_add=1.0",
            "intent_to_units_factor = AVG(units_forecast / weighted_intent_score); default 0.019",
            "intent_units_lift = ROUND(weighted_intent_score × intent_to_units_factor)",
            "units_historical = DemandForecast_Base.units_forecast at grain",
            "replacement_units_lift = SUM(replacement_score) for same run_id + grain (after RC)",
            "units_intent_adjusted = units_historical + intent_units_lift + replacement_units_lift",
            "confidence_score = MIN(1, (weighted_intent_score/1000) × (1 − guest_share))",
            "consumer_signals_applied = TRUE when Run 2 enabled",
        ],
        "outputs": [
            ("intent_units_lift", "demand_forecast_recommendation"),
            ("units_intent_adjusted", "demand_forecast_recommendation"),
            ("confidence_score", "demand_forecast_recommendation"),
            ("consumer_signals_applied", "demand_forecast_recommendation"),
            (
                "Recommendation header",
                "recommendation_id, baseline_forecast, adjusted_forecast, delta_units → recommendation",
            ),
            ("Week detail / drivers", "recommendation_traceability"),
        ],
    },
    {
        "name": "ReplacementCycleAgent",
        "role": (
            "Models repurchase likelihood windows by member segment × product. Identifies repeat "
            "buyers due to repurchase and translates that into a unit lift. Runs in parallel with "
            "DemandSensingAgent after Orchestrator. Separate signal path — does not read clickstream "
            "and does not replace intent scoring."
        ),
        "status": "TO-BE — glossary inputs only; no ConsumerAffinity or DemandForecast consumer tables",
        "inputs": [
            (
                "Orders",
                "member_id, product_cd, sku_id, transaction_dt_utc, sale_qty, "
                "transaction_type_desc, cancel_ind, enterprise_region",
            ),
            (
                "Member Hub",
                "member_id, membership_tier, clv_tier, region, last_purchase_dt",
            ),
            (
                "Product",
                "product_cd, category_cd, category_nm, subcategory_nm, sport_activity",
            ),
            ("Geo Region", "enterprise_region, region_canonical, client_region_cd"),
            ("DemandForecast_Base", "product_cd, region, week — routed scope only"),
        ],
        "input_note": (
            "Does not read: Clickstream Agg, Clickstream Base, ConsumerAffinity_Consumer, "
            "DemandForecast_Consumer. Sub-agent: CohortScoringSubAgent (BQML on inter-purchase intervals)."
        ),
        "computation": [
            "Filter Orders — 12m, SALES_ORDER, cancel_ind = 0, member_id not null",
            "Join Member Hub on member_id; Join Product on product_cd for category_cd",
            "last_purchase_dt = MAX(transaction_dt_utc) per member_id, product_cd",
            "category_cadence_days = MEDIAN(DATEDIFF between purchases) by category_cd",
            "repurchase_probability = CohortScoringSubAgent (BQML on inter-purchase intervals)",
            "replacement_due_count = COUNT(member_id) WHERE ABS(days_since_last − cadence) ≤ tolerance",
            "avg_repurchase_units = AVG(sale_qty) per member-product repurchase",
            "replacement_units_lift = ROUND(replacement_due_count × avg_repurchase_units) → replacement_score",
        ],
        "outputs": [
            ("repurchase_probability", "replacement_score"),
            ("replacement_due_count", "replacement_score"),
            ("avg_repurchase_units", "replacement_score"),
            (
                "replacement_units_lift",
                "replacement_score → blended into demand_forecast_recommendation by DS",
            ),
            ("replacement_driver", "recommendation_traceability (INS-05)"),
        ],
        "output_note": (
            "RC writes to replacement_score. DS reads replacement_units_lift and merges "
            "into units_intent_adjusted on demand_forecast_recommendation."
        ),
    },
    {
        "name": "InventoryPolicyAgent",
        "role": (
            "Applies business rules to combine demand sensing and replacement signals "
            "into safety stock recommendations. UC2 handoff."
        ),
        "status": "TO-BE (UC2 scope; included in Miro hourly loop)",
        "inputs": [
            (
                "demand_forecast_recommendation",
                "product_cd, sku_id, region_canonical, week, run_id, units_historical, "
                "units_intent_adjusted, intent_units_lift, replacement_units_lift, confidence_score",
            ),
            (
                "replacement_score",
                "product_cd, region_canonical, week, run_id, replacement_due_count, "
                "replacement_units_lift, repurchase_probability",
            ),
            ("DemandForecast_Base", "product_cd, region, week — baseline reference"),
            ("Inventory / WMS (TO-BE)", "On-hand, lead time, service level target"),
        ],
        "outputs": [
            ("safety_stock_units", "inventory_recommendation"),
            ("exception_priority_flag", "inventory_recommendation"),
            ("reorder_qty", "inventory_recommendation"),
            ("policy_rule_version, inputs_used", "inventory_recommendation"),
        ],
    },
    {
        "name": "TraceabilityAgent",
        "role": (
            "Wraps every recommendation with lineage — source tables, feature drivers, "
            "model version, and planner-readable reasoning."
        ),
        "status": "Partial — envelope built inside Demand Sensing pipeline; standalone agent TO-BE",
        "inputs": [
            (
                "demand_forecast_recommendation",
                "run_id, product_cd, region_canonical, week, intent_units_lift, "
                "replacement_units_lift, units_intent_adjusted, confidence_score",
            ),
            (
                "replacement_score",
                "run_id, replacement_due_count, replacement_units_lift, repurchase_probability",
            ),
            ("inventory_recommendation", "run_id, safety_stock_units, exception_priority_flag"),
            (
                "Source table refs",
                "Clickstream Agg, DemandForecast_Base, Orders, Member Hub, Product",
            ),
        ],
        "outputs": [
            ("signal, weight, source_table, detail", "recommendation_traceability"),
            ("Top drivers (feature contribution %)", "recommendation_traceability / INS-01…05"),
            ("agent_reasoning", "recommendation_traceability"),
            ("model_version", "recommendation_traceability"),
        ],
    },
    {
        "name": "PublisherAgent",
        "role": (
            "Publishes finalized recommendations to Pub/Sub and Insight API. "
            "Idempotent by recommendation_id. Does not apply forecasts."
        ),
        "status": "TO-BE",
        "inputs": [
            (
                "demand_forecast_recommendation",
                "run_id, product_cd, sku_id, region_canonical, week, units_historical, "
                "units_intent_adjusted, confidence_score",
            ),
            (
                "recommendation_traceability",
                "signal, weight, source_table, detail, model_version, agent_reasoning",
            ),
            (
                "inventory_recommendation",
                "run_id, safety_stock_units, exception_priority_flag (UC2 handoff)",
            ),
        ],
        "outputs": [
            (
                "recommendation_id, recommendation_type, geo_scope, sku, product_cd, "
                "baseline_forecast, adjusted_forecast, confidence, status",
                "recommendation",
            ),
            ("RecommendationEmitted event", "Pub/Sub topic"),
            ("Insight API row", "Planner insight feed"),
            ("status pending → approved / rejected / overridden", "recommendation"),
        ],
    },
    {
        "name": "InquiryAgent",
        "role": (
            "On-demand planner Q&A — e.g. what's driving demand for this SKU in PNW? "
            "Separate from the hourly loop."
        ),
        "status": "TO-BE",
        "inputs": [
            (
                "demand_forecast_recommendation",
                "product_cd, sku_id, region_canonical, week, units_historical, units_intent_adjusted, "
                "intent_units_lift, replacement_units_lift, confidence_score, consumer_signals_applied",
            ),
            (
                "replacement_score",
                "replacement_due_count, repurchase_probability, replacement_units_lift",
            ),
            (
                "recommendation",
                "recommendation_id, status, baseline_forecast, adjusted_forecast",
            ),
            (
                "recommendation_traceability",
                "signal, weight, source_table, detail, agent_reasoning",
            ),
            (
                "clickstream_agg",
                "search_count, pdp_view_count, cart_add_count, top_search_query, "
                "weighted_intent_score, member_intent_count, guest_intent_count",
            ),
        ],
        "outputs": [
            ("NL answer", "Plain-language explanation with trace refs"),
            ("Structured metrics", "Units, lift, confidence for requested scope"),
            ("Optional insight card", "Same format as push recommendations"),
        ],
    },
]

END_OF_RUN_TABLES = [
    (
        "signal_delta",
        "ForecastSignalBQAgent",
        "product_cd × region × week × delta_type",
        "delta_type, delta_magnitude, source_table, detected_at, region_canonical, prior_value, current_value",
    ),
    (
        "agent_run_history",
        "OrchestratorAgent",
        "per run_id",
        "run_id, run_status, gates_passed, scope_json, agents_executed, started_at, completed_at",
    ),
    (
        "demand_forecast_recommendation",
        "DemandSensingAgent (+ RC blend)",
        "product_cd × region × week × run_id",
        "units_historical, intent_units_lift, replacement_units_lift, units_intent_adjusted, "
        "confidence_score, consumer_signals_applied, recommendation_id",
    ),
    (
        "replacement_score",
        "ReplacementCycleAgent",
        "segment × product_cd × region × week × run_id",
        "repurchase_probability, replacement_due_count, avg_repurchase_units, "
        "replacement_units_lift, membership_tier",
    ),
    (
        "inventory_recommendation",
        "InventoryPolicyAgent",
        "sku_id × location × run_id",
        "safety_stock_units, exception_priority_flag, reorder_qty, policy_rule_version — UC2 handoff",
    ),
    (
        "recommendation",
        "PublisherAgent",
        "per recommendation_id",
        "recommendation_id, run_id, product_cd, sku_id, geo_scope, baseline_forecast, "
        "adjusted_forecast, confidence, status (pending)",
    ),
    (
        "recommendation_traceability",
        "TraceabilityAgent",
        "per recommendation_id × signal",
        "signal, weight, source_table, detail, model_version, agent_reasoning",
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    shading = cell._tc.get_or_add_tcPr()
    shd = shading.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd",
        {
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill": fill,
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val": "clear",
        },
    )
    shading.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "1B2A4A")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("UC1 Agent Specs")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    for line in [
        "Dataset: demandsensinglayer.dsl_dataset",
        "Grain: product_cd × region × week",
        "Phase 1: Recommend-only — planner approves before units_final",
    ]:
        r = meta.add_run(line + "\n")
        r.font.size = Pt(10)
        r.font.color.rgb = GREY

    for idx, agent in enumerate(AGENTS, start=1):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(f"{idx}. {agent['name']}")
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = TEAL

        role = doc.add_paragraph()
        role.paragraph_format.space_after = Pt(4)
        r = role.add_run("Role: ")
        r.bold = True
        r.font.size = Pt(11)
        r = role.add_run(agent["role"])
        r.font.size = Pt(11)

        status = doc.add_paragraph()
        status.paragraph_format.space_after = Pt(8)
        r = status.add_run("Status: ")
        r.bold = True
        r.font.size = Pt(10)
        r = status.add_run(agent["status"])
        r.font.size = Pt(10)
        r.font.color.rgb = GREY

        lbl = doc.add_paragraph()
        lbl.paragraph_format.space_after = Pt(4)
        r = lbl.add_run("Input tables and fields")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = NAVY

        add_table(doc, ["Table", "Fields"], agent["inputs"])

        if agent.get("input_note"):
            note = doc.add_paragraph()
            note.paragraph_format.space_after = Pt(8)
            r = note.add_run(agent["input_note"])
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = GREY

        if agent.get("computation"):
            lbl = doc.add_paragraph()
            lbl.paragraph_format.space_after = Pt(4)
            r = lbl.add_run("Runtime computation (each rerun)")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = NAVY
            for step in agent["computation"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(step)
                r.font.size = Pt(10)

        lbl = doc.add_paragraph()
        lbl.paragraph_format.space_after = Pt(4)
        r = lbl.add_run("Output")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = NAVY

        if agent.get("output_fields"):
            add_table(
                doc,
                ["Field", "Description", "Lands in"],
                agent["output_fields"],
            )
        else:
            out_headers = (
                ["Output", "Lands in"]
                if agent["name"]
                in {
                    "DemandSensingAgent",
                    "ReplacementCycleAgent",
                    "InventoryPolicyAgent",
                    "TraceabilityAgent",
                    "PublisherAgent",
                }
                else ["Output", "Description"]
            )
            add_table(doc, list(out_headers), agent["outputs"])

        if agent.get("does_not"):
            lbl = doc.add_paragraph()
            lbl.paragraph_format.space_after = Pt(4)
            r = lbl.add_run("Does NOT")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = NAVY
            for item in agent["does_not"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(item)
                r.font.size = Pt(10)

        if agent.get("output_note"):
            note = doc.add_paragraph()
            note.paragraph_format.space_after = Pt(8)
            r = note.add_run(agent["output_note"])
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = GREY

    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(20)
    h.paragraph_format.space_after = Pt(6)
    r = h.add_run("End-of-run tables")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = TEAL

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    r = intro.add_run(
        "Persisted at the end of each hourly agent loop. Input hubs (Orders, Member Hub, "
        "Product, Clickstream Agg, DemandForecast_Base) are pre-existing — not created per run. "
        "Planner HITL write-back (approval_audit + units_final) is separate — not end of hourly run."
    )
    r.font.size = Pt(11)

    add_table(
        doc,
        ["Table", "Written by", "Grain", "Key fields"],
        END_OF_RUN_TABLES,
    )

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(16)
    r = footer.add_run("UC1 · Forward Deploy Consulting · July 2026")
    r.font.size = Pt(9)
    r.font.color.rgb = GREY

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
