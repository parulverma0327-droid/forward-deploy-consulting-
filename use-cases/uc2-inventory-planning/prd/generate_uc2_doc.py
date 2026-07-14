"""Generate UC2 PRD docx — cover, index, and all sections with available content."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(
    os.path.dirname(__file__),
    "Retail AI- USE case 2.docx",
)

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
TEAL = RGBColor(0x1A, 0x6B, 0x5A)
GREY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.43)
    section.left_margin = Inches(0.69)
    section.right_margin = Inches(0.69)


def set_run(run, bold=False, size=11, colour=BLACK, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = colour


def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(text), bold=True, size=20, colour=NAVY)
    return p


def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(text), bold=True, size=14, colour=TEAL)
    return p


def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(text), bold=True, size=12, colour=NAVY)
    return p


def body(text, colour=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(text), size=11, colour=colour)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(text), size=11)
    return p


def numbered(text, level=0, bold=False):
    p = doc.add_paragraph(style="List Number")
    if level:
        p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    set_run(p.add_run(text), bold=bold, size=11)
    return p


def placeholder(text):
    body(text, colour=GREY)


def shade_header_row(table):
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1B2A4A")
        shd.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shd)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
            for para in table.rows[ri + 1].cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    shade_header_row(table)
    doc.add_paragraph()
    return table


def page_break():
    doc.add_page_break()


def cover():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    set_run(p.add_run("Retail AI — Use Case 2"), bold=True, size=28, colour=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    set_run(p.add_run("Inventory Planning and Allocation Loop"), italic=True, size=14, colour=TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    set_run(p.add_run("Forward Deploy Consulting · May 2026 · Confidential"), size=10, colour=GREY)

    for _ in range(17):
        doc.add_paragraph()

    doc.add_paragraph("Index")
    items = [
        ("Dictionary of Terms", 0, False),
        ("Use case and Problem statement", 0, True),
        ("Impact", 0, False),
        ("Impacted Domains & Systems", 0, True),
        ("As is and To Be process flows", 0, True),
        ("Blind spots/Friction Points", 0, True),
        ("Solution and Requirement", 0, True),
        ("In Scope", 1, True),
        ("Out of Scope", 1, True),
        ("Assumptions", 1, True),
        ("Implementation Approach", 0, False),
        ("Business Glossary + Data Standards", 0, True),
        ("Taxonomy  —  Excel — Tab 2", 0, False),
        ("Data Dictionary", 0, True),
        ("Ontology  —  Excel — Tab 4", 0, False),
        ("Semantic Layer  —  Excel — Tab 5", 0, False),
    ]
    for text, level, bold in items:
        numbered(text, level=level, bold=bold)


def section_dictionary():
    page_break()
    heading1("Dictionary of Terms")
    add_table(
        ["Term", "Description"],
        [
            ("Inventory Planning", "Monthly decisions on how much stock sits at each DC and store — after demand forecast and buy quantities are set. Distinct from weekly replenishment (UC3)."),
            ("Allocation", "The process of splitting available inventory across DCs and stores. Primary output is store/DC stock targets and transfer order quantities."),
            ("Allocation Engine", "Planning or ERP module that computes store-level stock targets from sales velocity, tier rules, and size curves. Often loosely connected to the approved demand forecast."),
            ("Inventory Hub", "System of record for on-hand, in-transit, and reserved inventory by SKU × location. Feeds allocation; often stale at store level."),
            ("Demand Forecast (input)", "Approved unit forecast from UC1 — SKU × region × week in Planning Data Hub. Allocation may reference it but as-is often re-derives need from POS velocity."),
            ("Sell-through / Velocity", "Units sold per week per store or region, computed from nightly POS batch. Primary as-is driver of allocation need."),
            ("Weeks of Cover", "On-hand units ÷ weekly sales rate. Common target metric (e.g. maintain 2–4 weeks at store)."),
            ("Store Tier (A/B/C)", "Merchandising grade of a store. Higher tiers receive larger allocation share."),
            ("Size Curve", "Historical % split of sales by size at a store or region. Applied to style-level allocation to get size-level qty."),
            ("Safety Stock", "Buffer units above expected demand. Set by planner rules or system defaults."),
            ("Transfer Order", "DC → store (or DC → DC) movement instruction. As-is often exported as a file and manually uploaded to WMS."),
            ("Replenishment (UC3)", "Weekly min/max restocking that maintains store shelves after allocation sets initial targets. Out of scope for this document."),
            ("FP – Friction Point", "A scoped break where a signal exists but fails to reach the system that needs it. UC2 FPs are numbered FP8 onwards."),
            ("Human in the loop", "AI agents propose allocation recommendations; allocation analysts review and approve before transfer orders are released."),
        ],
    )


def section_use_case():
    page_break()
    heading1("Use case and Problem statement")

    heading2("Use Case")
    body(
        "UC2 — Inventory Planning: Batch Signals, Blind Allocation. "
        "Once demand forecasting (UC1) produces an approved unit forecast, the business must decide "
        "where inventory sits in the network — which DCs hold stock, which stores receive initial "
        "allocations, and how regional pools are split. This runs on a monthly cadence at most "
        "retailers (weekly re-allocation in-season is common but still a position decision, not "
        "replenishment). Bigger companies usually separate inventory management (SAP, o9) from "
        "weekly replenishment; mid-market may combine them in one tool."
    )

    heading2("Problem Statement")
    body("The core problem: Demand signal and inventory position are disconnected from allocation decisions.")
    body(
        "Demand forecasting answers how many units. Inventory planning answers where those units sit. "
        "In as-is retail, allocation does not cleanly consume the approved forecast. It re-computes "
        "store need from historical sales velocity (nightly POS batch), tier rules, and size curves — "
        "while on-hand, in-transit, and DC constraints are often wrong, stale, or invisible."
    )
    body("The allocation signal that does exist is structurally limited:")
    bullet("Velocity-driven — store need = weeks of cover × last 4–8 weeks sales; enriched forecast from UC1 is optional, not default")
    bullet("Geography-blind — stock split by store tier or even %, not by where high-value customers or forward demand is concentrated")
    bullet("Position-stale — DC on-hand may be daily; store counts are often wrong or 24–48h+ late")
    bullet("Constraint-blind — transfer orders go out without DC capacity / pick backlog visibility")
    bullet("Manual execution — allocation output is often an Excel file manually uploaded to WMS")
    bullet("Consumer signals absent — CLV by geography, member concentration, and forward intent never enter allocation math")

    heading3("Why This Matters Now")
    body(
        "UC1 improves the demand number. If allocation still runs on blind, velocity-only rules, "
        "the right forecast still lands in the wrong stores. Markdown, stockouts, and DC backlog "
        "persist — the forecast accuracy gain never reaches the shelf."
    )


def section_impact():
    page_break()
    heading1("Impact")
    heading2("Business Impact (Estimated — $1–2B Retailer)")
    placeholder("[ Quantified impact model to be validated with client benchmarks ]")
    add_table(
        ["Impact Area", "Metric", "Estimated Value"],
        [
            ("Allocation accuracy", "10–20% reduction in regional misallocation", "$6–12M recovered margin / avoided markdown"),
            ("DC backlog / late transfers", "15–25% reduction in transfer cycle time", "$3–6M working capital release"),
            ("Member stockouts in priority geos", "20–30% reduction where CLV skew applied", "$4–8M recovered revenue"),
            ("Planner productivity", "30% reduction in manual allocation spreadsheet work", "$2–4M FTE redeployment"),
            ("Inventory distortion", "10–15% reduction in wrong-place overstock", "$8–15M markdown avoidance"),
            ("Total estimated Year 1 impact", "", "$23M–$45M"),
        ],
    )
    heading2("Operational Impact")
    bullet("Allocation analyst spends less time reconciling Excel exports with WMS reality")
    bullet("Transfer orders reflect DC capacity and in-transit position, not just theoretical need")
    bullet("Approved demand forecast from UC1 becomes an explicit allocation input, not a parallel spreadsheet")


def section_domains():
    page_break()
    heading1("Impacted Domains & Systems")

    heading2("Domains")
    heading3("1. Inventory & Allocation")
    body(
        "The primary domain. Decides stock placement across DCs and stores after buy/PO decisions. "
        "Consumes sales velocity, on-hand, in-transit, product master, and (weakly) demand forecast."
    )
    bullet("Store-level allocation and stock targets")
    bullet("DC → store transfer planning")
    bullet("Regional inventory positioning")
    bullet("Size-level size-curve splits")

    heading3("2. Planning (downstream of UC1)")
    body("Produces and approves DemandForecast (SKU × region × week). UC2 consumes units_final as an input to allocation — not rebuilt in UC2.")
    bullet("Approved demand forecast handoff")
    bullet("In-season forecast refresh feeding re-allocation")

    heading3("3. Product")
    body("SKU, style, size, colorway, season, flow dates — required to split allocation correctly.")
    bullet("Item master / UPC")
    bullet("Size curve definitions")
    bullet("Season and flow calendar")

    heading3("4. Sales Orders / POS")
    body("Nightly batch sales drive as-is velocity calculations. Not real-time for UC2 monthly planning; daily EOD is the floor.")
    bullet("Store POS transactions (EOD batch)")
    bullet(".com / OMS orders by ship-from location")

    heading3("5. Consumer & Member (input to allocation — missing as-is)")
    body("CLV by geography, member concentration, forward intent — should skew allocation; today stays in marketing stack.")
    bullet("Member sales by geography (from UC1 Member Sales View)")
    bullet("CLV tier by region (Data Science output)")

    heading2("Impacted Systems / Applications")
    add_table(
        ["System", "Domain", "What It Holds", "Role in UC2"],
        [
            ("WMS", "Inventory", "On-hand, in-transit, pick/ship execution", "Source of position; receives transfer orders"),
            ("ERP / SAP", "Finance + Inventory", "POs, receipts, inventory valuation", "Inventory management module; PO pipeline"),
            ("Planning System", "Planning", "Forecasts, allocation decisions", "Where allocation is computed (o9, Toolio, SAP IM)"),
            ("POS", "Store", "Transactions", "EOD batch → velocity for allocation"),
            ("OMS", ".com", "Orders, ship-from DC", "Online sell-through by fulfillment node"),
            ("Product Data Hub", "Product", "SKU, size, season", "Allocation grain and size curves"),
            ("Inventory Hub", "Inventory", "Aggregated OH + in-transit", "Position snapshot for allocation (often stale at store)"),
        ],
    )


def section_flows():
    page_break()
    heading1("As is and To Be process flows")

    heading2("As-Is: After Demand Forecast Is Approved")
    body("The following is the typical as-is chain once UC1 (or legacy planning) produces an approved demand forecast:")
    numbered("Demand forecast approved — SKU × region × week (Planning Data Hub or planning tool)")
    numbered("Buy / PO — buyers convert forecast + lead time to purchase orders (SAP/ERP). Separate process; not allocation.")
    numbered("Inbound — vendor → freight → DC receipt (WMS records on-hand at DC)")
    numbered("Inventory planning / allocation (~monthly) — compute store need from velocity + tier + size curve; compare to on-hand + in-transit; output transfer qty by store")
    numbered("Transfer execution — allocation file exported; analyst uploads to WMS (manual handoff)")
    numbered("Replenishment (~weekly, UC3) — min/max rules maintain store shelves against targets set by allocation")

    heading3("As-Is Allocation Logic (summary)")
    add_table(
        ["Step", "Logic", "Primary input"],
        [
            ("1. Compute need", "Target weeks of cover × weekly sales rate (last 4–8 weeks POS)", "POS batch"),
            ("2. Gap", "Need − (store on-hand + in-transit from DC)", "WMS / ERP (often stale)"),
            ("3. Split pool", "Pro-rata by tier, even %, or last-season share — not CLV/geo", "Merch rules"),
            ("4. Size break", "Apply historical size curve % to style qty", "POS history"),
            ("5. Override", "Analyst adjusts in Excel or planning UI", "Manual"),
            ("6. Output", "Transfer order file → WMS upload", "Allocation engine"),
        ],
    )

    heading3("As-Is: Inputs and Outputs")
    add_table(
        ["Direction", "Artifact", "Grain", "Notes"],
        [
            ("Input", "Approved demand forecast", "SKU × region × week", "Exists; allocation often ignores or loosely references"),
            ("Input", "Sales / sell-through", "SKU × store × week", "Nightly POS batch — primary driver"),
            ("Input", "On-hand + in-transit", "SKU × location", "DC better than store"),
            ("Input", "Open POs / pipeline", "SKU × DC", "What's already coming"),
            ("Input", "Product master + size curve", "SKU / style", "Size split"),
            ("Input", "Store tier, DC mapping", "Store / DC", "Who gets how much"),
            ("Output", "Allocation qty", "SKU × store", "Stock target or transfer qty"),
            ("Output", "Transfer order file", "SKU × DC → store", "Manual WMS upload"),
            ("Output", "Replenishment parameters", "Store min/max", "Feeds UC3 weekly runs"),
        ],
    )

    heading2("To-Be Process Flows")
    placeholder("[ To be detailed in next draft — Inventory Hub architecture, agentic allocation workflow, connection to UC1 DemandForecast and UC3 replenishment ]")


def section_fps():
    page_break()
    heading1("Blind spots / Friction Points")
    body("Seven friction points are in scope for UC2 (numbered FP8 onwards; UC1 covered FP1–FP7).")
    add_table(
        ["ID", "Description", "Type", "Priority", "Impact"],
        [
            (
                "FP8",
                "Demand Forecast Not Consumed by Allocation — Approved forecast exists in Planning Data Hub but allocation re-derives need from raw POS velocity. Enriched forecast from UC1 does not change store-level stock targets.",
                "Integration",
                "High",
                "Forecast accuracy investment does not reach the shelf. Right units, wrong locations.",
            ),
            (
                "FP9",
                "Allocation Geography-Blind — Stock split by store tier or even %, not by member concentration, CLV by geography, or regional forward demand. Node 2 (CRM segments → allocation) is missing.",
                "Integration",
                "High",
                "High-value geographies understocked; low-value stores overstocked.",
            ),
            (
                "FP10",
                "DC Capacity Blind Spot — Transfer orders sent without visibility into DC pick backlog, capacity, or receiving constraints. Orders pile up; replenishment arrives late anyway.",
                "Data + Integration",
                "High",
                "Transfer cycle time unpredictable; store stockouts despite DC having inventory.",
            ),
            (
                "FP11",
                "Allocation → WMS Manual Handoff — Allocation outputs a transfer file; a human uploads it to WMS. Delay, error-prone, no audit trail.",
                "Integration",
                "Medium",
                "Execution lag between plan and physical movement; reconciliation burden.",
            ),
            (
                "FP12",
                "Store On-Hand Stale or Wrong — Store inventory counts in WMS/POS do not match physical shelf. Allocation math uses bad position data.",
                "Data",
                "High",
                "Over-allocation to stores that cannot sell it; under-allocation where stock exists on paper only.",
            ),
            (
                "FP13",
                "Inventory-Triggered Markdown Not Connected — Aging stock in Inventory Hub does not trigger targeted markdown to the right customer segment. Promotions stay category-wide (Offer Irrelevance).",
                "Integration",
                "Medium",
                "Markdown margin loss; inventory ages in wrong locations.",
            ),
            (
                "FP14",
                "Loyalty Data Not Operationalised in Allocation — Member/loyalty insights used for dashboards only; not routed to allocation priorities or store stock skew.",
                "Integration",
                "Medium",
                "Member experience investments do not improve stock placement for members.",
            ),
        ],
    )


def section_solution():
    page_break()
    heading1("Solution and Requirement")
    placeholder("[ Placeholder — Inventory Hub architecture, Allocation Agent workflow, connection to UC1 PDH and UC3 replenishment parameters ]")
    body(
        "Direction (not yet detailed): Part 1 — Inventory Hub gold contract (position, in-transit, "
        "DC constraints, allocation targets). Part 2 — Inventory Planning Agent proposes store/DC "
        "reallocation from approved DemandForecast + position + member/geo signals; human analyst "
        "approves before transfer order release."
    )

    heading2("In Scope")
    bullet("Use Case 2 only: monthly inventory planning and allocation — stock levels across DCs and stores")
    bullet("As-is allocation process documentation and friction points FP8–FP14")
    bullet("Inventory Hub data definitions at aggregated / hub level (on-hand, in-transit, transfer orders)")
    bullet("Consumption of UC1 approved DemandForecast as explicit allocation input")
    bullet("Allocation analyst human-in-the-loop workflow (propose → review → approve transfer)")
    bullet("Dependency on UC1 identity and member-attributed demand for geography-skewed allocation (FP9)")
    bullet("Mid-market retail target ($100M–$500M); Nike-scale as reference architecture")

    heading2("Out of Scope")
    bullet("UC1 — Demand forecasting and consumer loop connection (separate document)")
    bullet("UC3 — Weekly replenishment, min/max execution, near-real-time store restocking")
    bullet("Buy plan / PO placement and vendor negotiation")
    bullet("Costing loop (landed cost, PLM cost negotiation) — future, not numbered UC")
    bullet("PLM feedback (returns → line plan) — future scope")
    bullet("Full WMS / ERP build — client provides hub-level inventory position via adapter")
    bullet("Real-time POS streaming for monthly allocation (daily EOD sufficient; UC3 needs fresher)")
    bullet("Custom allocation UI in Phase 1 — recommendations via existing planning tools or export")

    heading2("Assumptions")
    add_table(
        ["#", "Assumption"],
        [
            ("A1", "UC1 DemandForecast (units_final) is available at SKU × region × week before allocation runs."),
            ("A2", "Client has WMS or ERP inventory module with DC-level on-hand at daily EOD minimum."),
            ("A3", "Store-level on-hand exists but may be unreliable — data quality work expected."),
            ("A4", "Allocation runs monthly (or ad-hoc in-season re-allocation); weekly execution is UC3."),
            ("A5", "Store tier (A/B/C) and DC-to-store mapping exist in client master data."),
            ("A6", "Transfer orders today exit via file or API to WMS — SenseAct proposes; client executes."),
            ("A7", "Member / CLV by geography available from UC1 pipeline for geography-skewed allocation."),
            ("A8", "Allocation analyst role exists and will participate in human-in-the-loop review."),
        ],
    )


def section_implementation():
    page_break()
    heading1("Implementation Approach")
    placeholder("[ Placeholder — phased delivery after UC1; Inventory Hub build, Allocation Agent, WMS handoff adapter ]")
    add_table(
        ["Phase", "Focus", "Deliverable"],
        [
            ("Phase 0", "Prerequisite", "UC1 DemandForecast + Member Sales View live in PDH"),
            ("Phase 1", "Data foundation", "Inventory Hub gold contract — OH, in-transit, DC constraints"),
            ("Phase 2", "Agent", "Inventory Planning Agent — propose allocation adjustments"),
            ("Phase 3", "Execution", "Transfer order export / WMS adapter; analyst approval audit trail"),
        ],
    )

    heading1("Business Glossary + Data Standards")
    placeholder("[ Companion Excel workbook — Taxonomy Tab 2, Data Dictionary, Ontology Tab 4, Semantic Layer Tab 5 ]")
    bullet("Taxonomy — Excel — Tab 2")
    bullet("Data Dictionary — inventory hub entities: InventoryPosition, TransferOrder, AllocationTarget")
    bullet("Ontology — Excel — Tab 4")
    bullet("Semantic Layer — Excel — Tab 5")


def main():
    cover()
    section_dictionary()
    section_use_case()
    section_impact()
    section_domains()
    section_flows()
    section_fps()
    section_solution()
    section_implementation()
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
